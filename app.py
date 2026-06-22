"""
CVup — Resume Tailoring Tool
============================

Workflow:
    1. Upload your CV (PDF or DOCX). The text is extracted automatically.
    2. Paste the target job advertisement into a separate text box.
    3. Click "Analyze & Recommend" — Gemini compares your CV against the job
       ad and returns actionable advice.
    4. Accept or reject each piece of advice individually.
    5. Click "Regenerate CV" — Gemini rewrites your CV applying only the
       accepted advice, and you download it back in the SAME file format you
       uploaded (PDF -> PDF, DOCX -> DOCX).

The Gemini API key is read securely from `st.secrets["GEMINI_API_KEY"]`.
"""

import io
import json
import os
import re

import streamlit as st

# --- Optional third-party libraries (imported defensively) ---------------- #
try:
    from google import genai
except ImportError:
    genai = None

try:
    from groq import Groq
except ImportError:
    Groq = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    Document = None


# --------------------------------------------------------------------------- #
# Configuration
# --------------------------------------------------------------------------- #
# Default model: gemini-2.5-flash has a far more generous free-tier quota than
# the newest "…-latest" alias (which maps to gemini-3.5-flash, only ~20/day).
# For even more free headroom use "gemini-2.5-flash-lite".
# Override in .streamlit/secrets.toml with:  GEMINI_MODEL = "gemini-2.5-flash-lite"
# Gemini models tried in order — quality first, then ever-larger free-tier
# quotas, so the chain effectively SUMS each model's separate daily quota
# (~20 + 20 + 500 + 1500 ≈ 2040 requests/day) before Groq even gets involved.
#   gemini-3.5-flash       best quality (20/day)
#   gemini-2.5-flash       strong       (20/day)
#   gemini-3.1-flash-lite  good + big   (500/day)  ← the real workhorse buffer
#   gemma-4-26b-a4b-it     solid + huge (1500/day, unlimited tokens)
# If every Gemini model errors, we fall back to Groq (see below).
# Override the whole chain with a single model via secrets: GEMINI_MODEL = "...".
DEFAULT_MODEL_CHAIN = [
    "gemini-3.5-flash",
    "gemini-2.5-flash",
    "gemini-3.1-flash-lite",
    "gemma-4-26b-a4b-it",
]
# Final fallback provider: if all Gemini models error (quota/429/etc.).
# Override in secrets with GROQ_MODEL. llama-3.3-70b follows instructions well.
DEFAULT_GROQ_MODEL = "llama-3.3-70b-versatile"
APP_TITLE = "Career Magic 🐶🐱"  # browser tab title (name added dynamically in-page)

st.set_page_config(
    page_title=APP_TITLE,
    page_icon="🐶",
    layout="wide",
    initial_sidebar_state="expanded",
)


# --------------------------------------------------------------------------- #
# Theme — warm, colourful, corgi & cat vibes for Dorka 🐾
# --------------------------------------------------------------------------- #
def inject_theme() -> None:
    """Custom CSS for a cheerful, encouraging atmosphere."""
    st.markdown(
        """
        <style>
        /* Warm sunrise gradient background */
        .stApp {
            background: linear-gradient(160deg, #fff7ed 0%, #ffe9f0 45%, #fef3c7 100%);
        }
        section[data-testid="stSidebar"] {
            background: linear-gradient(180deg, #ffe4e6 0%, #ffedd5 100%);
        }
        /* Headings in a warm, friendly tone */
        h1, h2, h3 { color: #b45309 !important; }
        /* Rounded, cozy buttons */
        .stButton > button, .stDownloadButton > button {
            border-radius: 999px;
            border: 2px solid #fdba74;
            background: #fff;
            color: #c2410c;
            font-weight: 700;
            transition: transform .08s ease, box-shadow .12s ease;
        }
        .stButton > button:hover, .stDownloadButton > button:hover {
            transform: translateY(-1px);
            box-shadow: 0 4px 14px rgba(251,146,60,.35);
            border-color: #fb923c;
            color: #9a3412;
        }
        /* Primary buttons get the warm fill */
        .stButton > button[kind="primary"],
        .stDownloadButton > button[kind="primary"] {
            background: linear-gradient(90deg, #fb923c 0%, #f472b6 100%);
            color: #fff;
            border: none;
        }
        /* Cozy rounded cards */
        div[data-testid="stExpander"], div[data-testid="stForm"],
        div[data-testid="stVerticalBlockBorderWrapper"] {
            border-radius: 18px !important;
        }
        /* Metrics pop a little */
        div[data-testid="stMetric"] {
            background: rgba(255,255,255,.55);
            border-radius: 16px;
            padding: .5rem .25rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def show_corgi(width: int = 220) -> None:
    """
    Show Dorka's corgi photo if it's present in assets/, otherwise a cheerful
    emoji banner so the page still feels complete.
    """
    for ext in ("png", "jpg", "jpeg", "webp"):
        path = os.path.join(ASSETS_DIR, f"corgi.{ext}")
        if os.path.exists(path):
            st.image(path, width=width, caption="Your cheer-squad captain 🐶")
            return
    st.markdown(
        "<div style='font-size:3rem;text-align:center;line-height:1'>"
        "🐶🐾🐱</div>"
        "<p style='text-align:center;color:#b45309;font-size:.85rem'>"
        "Pop a corgi photo at <code>assets/corgi.png</code> to see it here 💛</p>",
        unsafe_allow_html=True,
    )


def show_corgi_in_sidebar() -> None:
    """Show Dorka's corgi at the top of the sidebar (photo if available)."""
    for ext in ("png", "jpg", "jpeg", "webp"):
        path = os.path.join(ASSETS_DIR, f"corgi.{ext}")
        if os.path.exists(path):
            st.sidebar.image(path, use_container_width=True)
            return
    st.sidebar.markdown(
        "<div style='font-size:2.4rem;text-align:center'>🐶</div>",
        unsafe_allow_html=True,
    )


ASSETS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")


# --------------------------------------------------------------------------- #
# Viewer identity (personalisation)
# --------------------------------------------------------------------------- #
# When the signed-in Streamlit viewer is the owner, the app greets her by name;
# everyone else is greeted by their own username.
OWNER_EMAIL = "sallosdorka@gmail.com"
OWNER_NAME = "Dorka"


def _viewer_email():
    """Best-effort signed-in viewer email (Streamlit Community Cloud / OIDC)."""
    for getter in (
        lambda: st.user.email,
        lambda: st.user.get("email"),
        lambda: st.experimental_user.email,
    ):
        try:
            email = getter()
            if email:
                return str(email)
        except Exception:
            continue
    return None


def resolve_viewer_name():
    """
    Decide who's viewing. Priority:
      1. A name the visitor typed in the sidebar (works on public apps).
      2. The signed-in viewer's email (private app or st.login) → owner / username.
      3. Unknown → None (greeted with a neutral fallback).
    """
    manual = (st.session_state.get("manual_name") or "").strip()
    if manual:
        return manual
    email = _viewer_email()
    if email and email.strip().lower() == OWNER_EMAIL:
        return OWNER_NAME
    if email:
        local = email.split("@")[0].replace(".", " ").replace("_", " ").strip()
        return local.title() if local else None
    return None


def who() -> str:
    """Display name for greetings; friendly fallback when unknown."""
    return st.session_state.get("_viewer_name") or "there"


def whos() -> str:
    """Possessive form, e.g. 'Dorka's' or 'Your'."""
    n = st.session_state.get("_viewer_name")
    return f"{n}'s" if n else "Your"


# --------------------------------------------------------------------------- #
# Session state
# --------------------------------------------------------------------------- #
def init_state() -> None:
    st.session_state.setdefault("cv_text", "")        # extracted CV text
    st.session_state.setdefault("cv_bytes", b"")       # original file bytes
    st.session_state.setdefault("cv_filename", "")     # original file name
    st.session_state.setdefault("cv_format", "")       # 'pdf' | 'docx' | 'txt'
    st.session_state.setdefault("job_ad", "")          # pasted job advert
    st.session_state.setdefault("company", "")          # target company (optional)
    # --- Recruiter & ATS analysis results --- #
    st.session_state.setdefault("match_score", None)     # int 0-100
    st.session_state.setdefault("missing_keywords", [])  # top 5
    st.session_state.setdefault("red_flags", [])         # 3 quick-spot issues
    # advice: actionable XYZ-rewrite / ATS items the user accepts or rejects.
    # each: {id, kind, title, original, suggestion, detail, status}
    st.session_state.setdefault("advice", [])
    st.session_state.setdefault("analysis_done", False)
    # --- Generated documents --- #
    st.session_state.setdefault("new_cv_text", "")     # finalized resume preview
    st.session_state.setdefault("new_cv_data", b"")    # resume bytes to download
    st.session_state.setdefault("new_cv_mime", "")
    st.session_state.setdefault("new_cv_ext", "")
    st.session_state.setdefault("cover_letter", "")     # tailored cover letter text
    st.session_state.setdefault("generation_done", False)
    # --- Interview prep --- #
    st.session_state.setdefault("interview_prep", {})    # dict of tip lists
    st.session_state.setdefault("interview_prep_done", False)


init_state()


# --------------------------------------------------------------------------- #
# Gemini helpers
# --------------------------------------------------------------------------- #
def get_api_key():
    try:
        return st.secrets["GEMINI_API_KEY"]
    except (KeyError, FileNotFoundError):
        return None


def get_client():
    """
    Return a cached google-genai Client, or None (after showing a friendly
    error) if the SDK or the API key is unavailable.
    """
    if genai is None:
        st.error(
            "The `google-genai` package is not installed. Install it with:"
            "\n\n```bash\npip install google-genai\n```"
        )
        return None

    api_key = get_api_key()
    if not api_key:
        st.error(
            f"🔑🐶 **The corgis can't find the magic key, {who()}!**\n\n"
            "This app needs a `GEMINI_API_KEY`. Add it to your Streamlit "
            "secrets:\n\n"
            "1. Create `.streamlit/secrets.toml` in this folder (or use the "
            "**Secrets** manager on Streamlit Community Cloud).\n"
            "2. Add this line:\n\n"
            "```toml\nGEMINI_API_KEY = \"your-api-key-here\"\n```\n\n"
            "3. Restart the app. Get a key at "
            "[Google AI Studio](https://aistudio.google.com/app/apikey)."
        )
        return None

    # Cache the client across reruns (keyed by api_key so a changed key rebuilds).
    cached = st.session_state.get("_gemini_client")
    if cached is None or st.session_state.get("_gemini_key") != api_key:
        st.session_state["_gemini_client"] = genai.Client(api_key=api_key)
        st.session_state["_gemini_key"] = api_key
    return st.session_state["_gemini_client"]


def get_model_chain() -> list:
    """
    Ordered list of Gemini models to try. A single GEMINI_MODEL in secrets
    overrides the whole chain; otherwise the default chain is used.
    """
    try:
        override = st.secrets.get("GEMINI_MODEL")
    except (FileNotFoundError, KeyError):
        override = None
    return [override] if override else list(DEFAULT_MODEL_CHAIN)


def get_groq_model() -> str:
    """Groq model id, overridable via st.secrets['GROQ_MODEL']."""
    try:
        return st.secrets.get("GROQ_MODEL", DEFAULT_GROQ_MODEL)
    except (FileNotFoundError, KeyError):
        return DEFAULT_GROQ_MODEL


def get_groq_client():
    """Return a cached Groq client, or None if the SDK/key are unavailable."""
    if Groq is None:
        return None
    try:
        key = st.secrets["GROQ_API_KEY"]
    except (KeyError, FileNotFoundError):
        return None
    if not key:
        return None
    cached = st.session_state.get("_groq_client")
    if cached is None or st.session_state.get("_groq_key") != key:
        st.session_state["_groq_client"] = Groq(api_key=key)
        st.session_state["_groq_key"] = key
    return st.session_state["_groq_client"]


def _is_quota_error(exc: Exception) -> bool:
    msg = str(exc)
    return "429" in msg or "RESOURCE_EXHAUSTED" in msg


class _Result:
    """Minimal stand-in exposing `.text`, so callers stay provider-agnostic."""

    def __init__(self, text: str):
        self.text = text


def _groq_call(groq_client, prompt: str) -> str:
    """One Groq (OpenAI-compatible) chat completion; returns the text."""
    resp = groq_client.chat.completions.create(
        model=get_groq_model(),
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
    )
    return resp.choices[0].message.content or ""


def _toast(message: str, icon: str = "🦙") -> None:
    """Non-blocking notification; ignored outside a Streamlit run context."""
    try:
        st.toast(message, icon=icon)
    except Exception:
        pass


def _generate(client, prompt: str):
    """
    Single entry point for all LLM calls. Tries each Gemini model in order
    (newest first, then higher-quota), and if every Gemini model errors, falls
    back to Groq. Returns an object exposing `.text`.
    """
    last_exc = None
    for idx, model in enumerate(get_model_chain()):
        try:
            resp = client.models.generate_content(model=model, contents=prompt)
            if idx > 0:  # a non-primary model answered
                _toast(f"Gemini fell back to {model} 🐾", icon="🐶")
            return resp
        except Exception as exc:  # noqa: BLE001 - try the next provider/model
            last_exc = exc
            continue

    # Every Gemini model failed → Groq fallback.
    groq_client = get_groq_client()
    if groq_client is not None:
        try:
            text = _groq_call(groq_client, prompt)
            _toast("Gemini's having a nap — the Groq llama took over! 🦙")
            return _Result(text)
        except Exception:
            raise last_exc  # surface the original Gemini error
    raise last_exc


def show_api_error(exc: Exception, what: str) -> None:
    """Render a friendly, quota-aware error message for Dorka."""
    backup = " (and the Groq llama backup couldn't catch it either 🦙)" if \
        get_groq_client() is not None else ""
    if _is_quota_error(exc):
        st.error(
            f"🐾 Phew, {who()} — we hit Gemini's request limit while {what}{backup}. "
            f"The current model gives only a small number of free requests, and "
            f"we've used them up for now."
        )
        st.info(
            "**Two ways to fix it:**\n\n"
            "1. 🐶 **Quick:** switch to a lighter model with a bigger free quota. "
            "Add this to `.streamlit/secrets.toml` and restart:\n\n"
            "```toml\nGEMINI_MODEL = \"gemini-2.5-flash-lite\"\n```\n\n"
            "2. 💳 **For real Pro limits:** enable billing on your API project at "
            "[Google AI Studio → Billing](https://aistudio.google.com/app/billing). "
            "Note: a *consumer* Gemini Pro / Google One subscription does **not** "
            "raise API limits — billing on the API project is what counts.\n\n"
            "_(Tip: a full run uses ~3 requests — one to analyse, one for the "
            "résumé, one for the cover letter.)_"
        )
    else:
        st.error(
            f"🐱 Oh no, a little hiccup while {what}: {exc}\n\n"
            f"Don't worry, {who()} — even the best corgis trip sometimes. "
            "Give it another try! 💛"
        )


def _extract_json(text: str):
    """Best-effort extraction of JSON from Gemini output (may be fenced)."""
    fenced = re.search(r"```(?:json)?\s*(.*?)```", text, re.DOTALL)
    if fenced:
        text = fenced.group(1)
    match = re.search(r"(\[.*\]|\{.*\})", text, re.DOTALL)
    candidate = match.group(1) if match else text
    return json.loads(candidate)


def analyze_cv(client, cv_text: str, job_ad: str) -> dict:
    """
    Run a 3-step Recruiter & ATS analysis comparing the resume to the job ad.

    Returns a dict:
        {
          "match_score": int,                 # 0-100
          "missing_keywords": [str, ...],      # top 5
          "red_flags": [str, ...],             # 3 quick-spot issues
          "recommendations": [                  # actionable, accept/reject items
            {"kind": "xyz", "title", "original", "suggestion", "detail"},
            {"kind": "ats", "title", "original", "suggestion", "detail"}
          ]
        }
    """
    prompt = f"""You are a senior technical recruiter and an ATS (Applicant
Tracking System) parser combined. Analyse the candidate's RESUME against the JOB
DESCRIPTION using this strict framework and be brutally honest but constructive.

Produce:
1. MATCH SCORE: an integer 0-100 estimating how well the resume fits the role.
2. MISSING KEYWORDS: the top 5 important keywords/skills from the job description
   that are absent or under-represented in the resume.
3. RED FLAGS: exactly 3 things a hiring manager would notice in under 10 seconds
   that hurt this candidate (e.g. vague bullets, no metrics, job-hopping,
   missing must-have, walls of text, irrelevant focus).
4. REWRITE RECOMMENDATIONS (Google XYZ formula): pick specific WEAK bullet points
   from the resume and rewrite them as "Accomplished [X] as measured by [Y], by
   doing [Z]". Quote the real original bullet in "original" and put the rewritten
   version in "suggestion". Stay truthful — if a metric is unknown, use a
   clearly placeholder like "[X]%" the candidate can fill in.
5. ATS OPTIMIZATION: identify sections an ATS or a tired hiring manager would
   skip or skim, and provide a "scroll-stopping" rewrite. Put what is being
   skipped in "original" and the punchier rewrite in "suggestion".

Return ONLY valid JSON in EXACTLY this shape (no prose, no markdown fences):
{{
  "match_score": <int>,
  "missing_keywords": ["", "", "", "", ""],
  "red_flags": ["", "", ""],
  "recommendations": [
    {{"kind": "xyz", "title": "", "original": "", "suggestion": "", "detail": ""}},
    {{"kind": "ats", "title": "", "original": "", "suggestion": "", "detail": ""}}
  ]
}}

Rules for "recommendations":
- Provide 4 to 7 items, mixing "xyz" (bullet rewrites) and "ats" (section
  optimizations).
- "title": short imperative summary (max ~12 words).
- "detail": 1-2 sentences on why this change helps with the recruiter/ATS.

=== RESUME ===
{cv_text}

=== JOB DESCRIPTION ===
{job_ad}
"""
    response = _generate(client, prompt)
    data = _extract_json(response.text)
    if not isinstance(data, dict):
        raise ValueError("Analysis response was not a JSON object.")

    # --- Normalise score --- #
    try:
        score = int(round(float(data.get("match_score", 0))))
    except (TypeError, ValueError):
        score = 0
    score = max(0, min(100, score))

    keywords = [str(k).strip() for k in (data.get("missing_keywords") or [])
                if str(k).strip()][:5]
    red_flags = [str(r).strip() for r in (data.get("red_flags") or [])
                 if str(r).strip()][:3]

    recommendations = []
    for idx, item in enumerate(data.get("recommendations") or []):
        if not isinstance(item, dict):
            continue
        kind = str(item.get("kind", "")).strip().lower()
        kind = "ats" if kind == "ats" else "xyz"
        recommendations.append(
            {
                "id": idx,
                "kind": kind,
                "title": str(item.get("title", "")).strip()
                or f"Suggestion {idx + 1}",
                "original": str(item.get("original", "")).strip(),
                "suggestion": str(item.get("suggestion", "")).strip(),
                "detail": str(item.get("detail", "")).strip(),
                "status": "pending",
            }
        )

    return {
        "match_score": score,
        "missing_keywords": keywords,
        "red_flags": red_flags,
        "recommendations": recommendations,
    }


def _accepted_instructions(accepted: list) -> str:
    """Format accepted recommendations (XYZ rewrites / ATS fixes) for prompts."""
    if not accepted:
        return "(No specific changes accepted; perform a light, sensible tailoring.)"
    lines = []
    for a in accepted:
        line = f"- {a.get('title', '').strip()}"
        original = a.get("original", "").strip()
        suggestion = a.get("suggestion", "").strip()
        detail = a.get("detail", "").strip()
        if original:
            line += f'\n    Replace this weak text: "{original}"'
        if suggestion:
            line += f'\n    With this stronger version: "{suggestion}"'
        if detail and not suggestion:
            line += f": {detail}"
        lines.append(line)
    return "\n".join(lines)


def _company_line(company: str) -> str:
    company = (company or "").strip()
    if company:
        return f"The target company is: {company}. Tailor wording to fit it."
    return "No specific company was given — keep it broadly professional."


def regenerate_cv(
    client, cv_text: str, job_ad: str, accepted: list, company: str = ""
) -> str:
    """Rewrite the CV applying only the accepted advice. Returns plain text."""
    accepted_text = _accepted_instructions(accepted)

    prompt = f"""You are an expert CV writer.

Rewrite the candidate's CV so it is tailored to the JOB ADVERTISEMENT, applying
ONLY the APPROVED CHANGES below. {_company_line(company)} Keep all facts truthful
— do NOT invent employers, degrees, dates, or experience not present in the
original CV; you may rephrase, reorder, and emphasise.

Preserve a clean, professional CV structure (clear section headings such as
SUMMARY, EXPERIENCE, EDUCATION, SKILLS; bullet points starting with "- ").
Return ONLY the full CV as plain text — no commentary, no markdown fences.

=== APPROVED CHANGES ===
{accepted_text}

=== ORIGINAL CV ===
{cv_text}

=== JOB ADVERTISEMENT ===
{job_ad}
"""
    response = _generate(client, prompt)
    return (response.text or "").strip()


def generate_cover_letter(
    client, finalized_resume_text: str, job_ad: str, accepted: list,
    company: str = "",
) -> str:
    """
    Write a tailored cover letter for the role whose tone matches the finalized
    (already optimized) resume. Returns plain text.
    """
    accepted_text = _accepted_instructions(accepted)

    prompt = f"""You are an expert career writer crafting a cover letter.

Write a tailored, compelling cover letter for the candidate applying to the role
in the JOB DESCRIPTION. {_company_line(company)} It must be fully customised to
THIS role and company, and match the tone and strengths of the FINALIZED RESUME
below.

Requirements:
- 3-4 short paragraphs, about 250-350 words total.
- Open with a hook that connects the candidate to this specific role/company.
- Weave in the most relevant, optimized achievements from the finalized resume
  (reflecting the approved improvements) — do NOT invent new facts.
- Mirror the keywords and priorities of the job description.
- Confident and warm, never generic or robotic.
- If the candidate's name appears in the resume, sign off with it; otherwise end
  with a simple "Sincerely,". Use "[Hiring Manager]" if no name is known.
- Return ONLY the cover letter as plain text — no commentary, no markdown fences.

=== APPROVED IMPROVEMENTS (already applied to the resume) ===
{accepted_text}

=== FINALIZED RESUME ===
{finalized_resume_text}

=== JOB DESCRIPTION ===
{job_ad}
"""
    response = _generate(client, prompt)
    return (response.text or "").strip()


def generate_interview_prep(
    client, cv_text: str, job_ad: str, company: str = ""
) -> dict:
    """
    Produce concrete interview-prep material tailored to the candidate + role.
    Returns a dict with four lists.
    """
    prompt = f"""You are an experienced interview coach.

Using the candidate's CV and the JOB DESCRIPTION, prepare focused, practical
interview preparation. {_company_line(company)}

Return ONLY valid JSON (no commentary, no markdown fences) in this shape:
{{
  "likely_questions": ["", "", "", "", ""],
  "your_talking_points": ["", "", ""],
  "questions_to_ask": ["", "", ""],
  "quick_tips": ["", "", ""]
}}

Guidance:
- "likely_questions": 5 specific questions this candidate will probably get for
  THIS role (mix behavioural + role/tech specific). Tie them to gaps or
  strengths visible in the CV vs the job.
- "your_talking_points": 3 concrete, truthful strengths/stories from the CV the
  candidate should lead with, phrased as quick reminders.
- "questions_to_ask": 3 sharp questions for the candidate to ask the
  interviewer (about the role, team, or company).
- "quick_tips": 3 short, encouraging, actionable prep tips.

=== CV ===
{cv_text}

=== JOB DESCRIPTION ===
{job_ad}
"""
    response = _generate(client, prompt)
    data = _extract_json(response.text)
    if not isinstance(data, dict):
        raise ValueError("Interview prep response was not a JSON object.")
    keys = ("likely_questions", "your_talking_points", "questions_to_ask",
            "quick_tips")
    return {
        k: [str(x).strip() for x in (data.get(k) or []) if str(x).strip()]
        for k in keys
    }


def regenerate_cv_structured(
    client, cv_text: str, job_ad: str, accepted: list, company: str = ""
) -> dict:
    """
    Tailor the CV and return it as a STRUCTURED dict (plus a styling "theme")
    so it can be rendered into a polished, position-tailored modern layout.
    """
    accepted_text = _accepted_instructions(accepted)
    company = (company or "").strip()
    company_note = (
        f'The target company is "{company}". Choose "accent_hex" to evoke that '
        f"company's brand colour (or its industry's vibe) while staying tasteful "
        f"and readable on white."
        if company else
        "No company was given — use a confident, modern professional accent "
        "colour (a tasteful teal, indigo, or slate works well)."
    )

    prompt = f"""You are an expert CV writer and modern resume designer.

Rewrite the candidate's CV so it is tailored to the JOB ADVERTISEMENT, applying
ONLY the APPROVED CHANGES below, and return it as STRUCTURED JSON that can be
laid out as a clean, contemporary resume for THIS specific role.

RULES:
- Keep every fact truthful. Do NOT invent employers, degrees, dates, or
  experience not present in the original CV. You may rephrase, reorder,
  emphasise, and inject role-relevant keywords.
- "title" is a short professional headline tailored to the target role
  (e.g. "Backend Engineer · Python & Cloud"). Infer it from the candidate's
  real background and the job ad; keep it credible.
- Lead with the most relevant sections for this job.
- Make experience bullets impact-oriented and, where the original supports it,
  quantified. Keep them concise (one line each ideally).
- Use these section "type" values:
    "paragraph" (with "content"), "bullets" (with "bullets"),
    or "entries" (with "entries": role/org/dates/bullets).
- If the original CV lacks a name or contact details, use empty string / empty
  list — do NOT fabricate them.
- "theme.accent_hex" is a 6-digit hex colour (no #) for headings/accents.
  {company_note}

Return ONLY valid JSON matching this shape (no commentary, no markdown fences):
{{
  "name": "",
  "title": "",
  "contact": [],
  "theme": {{"accent_hex": "0F766E"}},
  "sections": [
    {{"heading": "", "type": "paragraph", "content": ""}},
    {{"heading": "", "type": "entries", "entries": [
        {{"role": "", "org": "", "dates": "", "bullets": [""]}}
    ]}},
    {{"heading": "", "type": "bullets", "bullets": [""]}}
  ]
}}

=== APPROVED CHANGES ===
{accepted_text}

=== ORIGINAL CV ===
{cv_text}

=== JOB ADVERTISEMENT ===
{job_ad}
"""
    response = _generate(client, prompt)
    data = _extract_json(response.text)
    if not isinstance(data, dict):
        raise ValueError("Structured CV response was not a JSON object.")
    return data


def _structured_to_text(data: dict) -> str:
    """Flatten a structured CV dict into readable plain text (for previews)."""
    lines = []
    if data.get("name"):
        lines.append(str(data["name"]))
    if data.get("title"):
        lines.append(str(data["title"]))
    contact = [str(c) for c in (data.get("contact") or []) if str(c).strip()]
    if contact:
        lines.append("  •  ".join(contact))
    for sec in (data.get("sections") or []):
        if not isinstance(sec, dict):
            continue
        lines.append("")
        if sec.get("heading"):
            lines.append(str(sec["heading"]).upper())
        for e in (sec.get("entries") or []):
            if not isinstance(e, dict):
                continue
            head = " — ".join(x for x in (e.get("role"), e.get("org")) if x)
            if e.get("dates"):
                head = f"{head}  ({e['dates']})" if head else str(e["dates"])
            if head:
                lines.append(head)
            for b in (e.get("bullets") or []):
                if str(b).strip():
                    lines.append(f"  - {b}")
        for b in (sec.get("bullets") or []):
            if str(b).strip():
                lines.append(f"  - {b}")
        if str(sec.get("content", "")).strip():
            lines.append(str(sec["content"]))
    return "\n".join(lines).strip()


# --------------------------------------------------------------------------- #
# DOCX in-place editing (preserves original styling)
# --------------------------------------------------------------------------- #
def _iter_docx_paragraphs(doc):
    """
    Yield every paragraph in a python-docx document, including those nested
    inside tables (CV templates often use tables for layout). Recurses into
    nested tables as well.
    """
    def walk_cell(cell):
        for p in cell.paragraphs:
            yield p
        for tbl in cell.tables:
            yield from walk_table(tbl)

    def walk_table(tbl):
        for row in tbl.rows:
            for cell in row.cells:
                yield from walk_cell(cell)

    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        yield from walk_table(tbl)


def _set_paragraph_text(paragraph, text: str) -> None:
    """
    Replace a paragraph's text while preserving its character formatting.
    The first run keeps its font/bold/size and receives the new text; any
    remaining runs are emptied. Paragraph-level style (heading, bullet,
    alignment, spacing) is untouched.
    """
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def regenerate_docx_in_place(
    client, original_bytes: bytes, job_ad: str, accepted: list, company: str = ""
) -> bytes:
    """
    Tailor a DOCX CV by rewriting only the TEXT of each existing paragraph,
    keeping the document's original styling and layout. Returns new DOCX bytes.
    """
    doc = Document(io.BytesIO(original_bytes))

    # Index only the paragraphs that actually contain text.
    paragraphs = list(_iter_docx_paragraphs(doc))
    indexed = [(i, p) for i, p in enumerate(paragraphs) if p.text.strip()]

    items = [{"i": i, "text": p.text} for i, p in indexed]

    accepted_text = _accepted_instructions(accepted)

    prompt = f"""You are an expert CV writer editing a CV in place.
{_company_line(company)}

You are given the CV as a JSON array of paragraph objects, each with an integer
"i" (index) and its current "text". Rewrite the "text" of each paragraph so the
CV is tailored to the JOB ADVERTISEMENT, applying ONLY the APPROVED CHANGES
below.

STRICT RULES:
- Return EXACTLY one object per input paragraph, with the SAME "i" values.
- Do NOT add, remove, split, merge, or reorder paragraphs.
- Keep each paragraph's length roughly similar so it still fits the layout.
- Keep all facts truthful — do NOT invent employers, degrees, dates, or
  experience. You may rephrase, re-emphasise, and inject relevant keywords.
- If a paragraph needs no change, return its text unchanged.
- Return ONLY a valid JSON array of objects: {{"i": <int>, "text": <string>}}.
  No commentary, no markdown fences.

=== APPROVED CHANGES ===
{accepted_text}

=== JOB ADVERTISEMENT ===
{job_ad}

=== CV PARAGRAPHS (JSON) ===
{json.dumps(items, ensure_ascii=False)}
"""

    response = _generate(client, prompt)
    data = _extract_json(response.text)

    # Map index -> new text, then write back preserving styling.
    new_by_index = {}
    for obj in data:
        if isinstance(obj, dict) and "i" in obj and "text" in obj:
            new_by_index[int(obj["i"])] = str(obj["text"])

    for i, paragraph in indexed:
        if i in new_by_index and new_by_index[i].strip():
            _set_paragraph_text(paragraph, new_by_index[i])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# File extraction
# --------------------------------------------------------------------------- #
def extract_text(uploaded) -> tuple:
    """
    Extract text from an uploaded file.
    Returns (text, fmt, raw_bytes) where fmt is 'pdf' | 'docx' | 'txt'.
    The raw bytes are kept so a DOCX can later be edited in place.
    Raises ValueError on unsupported/missing-dependency cases.
    """
    name = uploaded.name.lower()
    raw = uploaded.read()

    if name.endswith(".pdf"):
        if PdfReader is None:
            raise ValueError("`pypdf` is not installed (needed to read PDFs).")
        reader = PdfReader(io.BytesIO(raw))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text.strip(), "pdf", raw

    if name.endswith(".docx"):
        if Document is None:
            raise ValueError("`python-docx` is not installed (needed for DOCX).")
        doc = Document(io.BytesIO(raw))
        text = "\n".join(p.text for p in _iter_docx_paragraphs(doc))
        return text.strip(), "docx", raw

    if name.endswith((".txt", ".md")):
        return raw.decode("utf-8", errors="replace").strip(), "txt", raw

    raise ValueError(
        "Unsupported file type. Please upload a PDF, DOCX, or TXT file. "
        "(Legacy .doc is not supported — save it as .docx first.)"
    )


# --------------------------------------------------------------------------- #
# Modern DOCX generation
# --------------------------------------------------------------------------- #
DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
DEFAULT_ACCENT = (15, 118, 110)   # refined teal — modern, professional default
INK = RGBColor(0x22, 0x22, 0x22) if Document else None
MUTED = RGBColor(0x6B, 0x72, 0x80) if Document else None


def _parse_accent(theme) -> tuple:
    """Read an accent RGB tuple from a theme dict's hex; fall back to default."""
    if isinstance(theme, dict):
        raw = str(theme.get("accent_hex") or theme.get("accent") or "").strip()
        raw = raw.lstrip("#")
        if len(raw) == 6:
            try:
                return tuple(int(raw[i:i + 2], 16) for i in (0, 2, 4))
            except ValueError:
                pass
    return DEFAULT_ACCENT


def _hexstr(rgb: tuple) -> str:
    return "".join(f"{c:02X}" for c in rgb)


def _bottom_border(paragraph, color_hex: str, sz: int = 6, space: int = 6) -> None:
    """Add a thin bottom rule under a paragraph (for modern section dividers)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))       # eighths of a point
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _run(paragraph, text, *, bold=False, italic=False, size=None, color=None,
         font_name="Calibri", spacing=None):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = font_name
    if size is not None:
        r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = RGBColor(*color) if isinstance(color, tuple) else color
    if spacing is not None:
        # Character spacing in twentieths of a point (for airy small-caps look).
        rPr = r._r.get_or_add_rPr()
        el = OxmlElement("w:spacing")
        el.set(qn("w:val"), str(spacing))
        rPr.append(el)
    return r


def build_docx_styled(data: dict, accent: tuple = DEFAULT_ACCENT) -> bytes:
    """
    Render a structured CV dict into a clean, modern, ATS-friendly Word document:
    generous whitespace, a bold name, an accent headline, accent section rules,
    role + right-aligned dates, italic org lines, and tidy bullets.
    """
    accent_hex = _hexstr(accent)
    doc = Document()

    # Modern page margins.
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)
    usable = doc.sections[0].page_width - doc.sections[0].left_margin \
        - doc.sections[0].right_margin

    # Base ("Normal") style: modern sans, comfortable spacing.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.08

    def spacer(pt=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        _run(p, "", size=pt)

    # --- Header --- #
    name = str(data.get("name", "")).strip()
    if name:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        _run(p, name, bold=True, size=26, color=INK)

    title = str(data.get("title", "")).strip()
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, title, size=12.5, color=accent)

    contact = [str(c).strip() for c in (data.get("contact") or []) if str(c).strip()]
    if contact:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _run(p, "   •   ".join(contact), size=9, color=MUTED)
        _bottom_border(p, accent_hex, sz=6, space=8)

    def section_header(heading):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        _run(p, heading.upper(), bold=True, size=11, color=accent, spacing=24)
        _bottom_border(p, accent_hex, sz=4, space=4)

    def paragraph(text):
        p = doc.add_paragraph()
        _run(p, str(text), size=10.5, color=INK)

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        _run(p, str(text), size=10.5, color=INK)

    def entry(e):
        role = str(e.get("role", "")).strip()
        org = str(e.get("org", "")).strip()
        dates = str(e.get("dates", "")).strip()
        if role or dates:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.tab_stops.add_tab_stop(
                usable, WD_TAB_ALIGNMENT.RIGHT
            )
            _run(p, role, bold=True, size=10.5, color=INK)
            if dates:
                _run(p, "\t" + dates, size=9.5, color=MUTED)
        if org:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _run(p, org, italic=True, size=10, color=accent)
        for b in (e.get("bullets") or []):
            if str(b).strip():
                bullet(b)
        spacer(3)

    for sec in (data.get("sections") or []):
        if not isinstance(sec, dict):
            continue
        if sec.get("heading"):
            section_header(str(sec["heading"]))
        stype = sec.get("type", "")
        if stype == "entries" or sec.get("entries"):
            for e in (sec.get("entries") or []):
                if isinstance(e, dict):
                    entry(e)
        elif stype == "bullets" or sec.get("bullets"):
            for b in (sec.get("bullets") or []):
                if str(b).strip():
                    bullet(b)
        else:
            if str(sec.get("content", "")).strip():
                paragraph(sec["content"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_docx_plain(text: str, accent: tuple = DEFAULT_ACCENT) -> bytes:
    """Modern DOCX from plain CV text (fallback when structuring fails)."""
    accent_hex = _hexstr(accent)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.isupper() and len(stripped) <= 40:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            _run(p, stripped, bold=True, size=11, color=accent, spacing=24)
            _bottom_border(p, accent_hex, sz=4, space=4)
        elif stripped.startswith(("- ", "• ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _run(p, stripped[2:].strip(), size=10.5, color=INK)
        else:
            p = doc.add_paragraph()
            _run(p, stripped, size=10.5, color=INK)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_prep_docx(prep: dict, job_ad_title: str = "") -> bytes:
    """Build a tidy, OneNote-friendly DOCX of interview prep tips."""
    accent_hex = _hexstr(DEFAULT_ACCENT)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    p = doc.add_paragraph()
    _run(p, "Interview Prep", bold=True, size=20, color=INK)
    if job_ad_title:
        sub = doc.add_paragraph()
        _run(sub, job_ad_title, size=11, color=MUTED)

    labels = {
        "likely_questions": "❓ Questions you'll likely be asked",
        "your_talking_points": "⭐ Your talking points",
        "questions_to_ask": "🙋 Smart questions to ask them",
        "quick_tips": "✅ Quick tips",
    }
    for key, label in labels.items():
        items = [str(x).strip() for x in (prep.get(key) or []) if str(x).strip()]
        if not items:
            continue
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        _run(h, label, bold=True, size=13, color=DEFAULT_ACCENT)
        _bottom_border(h, accent_hex, sz=4, space=4)
        for it in items:
            b = doc.add_paragraph(style="List Bullet")
            _run(b, it, size=11, color=INK)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def prep_to_text(prep: dict, job_ad_title: str = "") -> str:
    """Flatten interview prep into copy-paste text (for OneNote / clipboard)."""
    lines = ["INTERVIEW PREP"]
    if job_ad_title:
        lines.append(job_ad_title)
    labels = {
        "likely_questions": "Questions you'll likely be asked",
        "your_talking_points": "Your talking points",
        "questions_to_ask": "Smart questions to ask them",
        "quick_tips": "Quick tips",
    }
    for key, label in labels.items():
        items = [str(x).strip() for x in (prep.get(key) or []) if str(x).strip()]
        if not items:
            continue
        lines.append("")
        lines.append(label)
        for it in items:
            lines.append(f"  - {it}")
    return "\n".join(lines).strip()


# --------------------------------------------------------------------------- #
# Sidebar — CV upload
# --------------------------------------------------------------------------- #
def render_sidebar() -> None:
    show_corgi_in_sidebar()
    st.sidebar.title(f"🐾 {whos()} CV Corner")
    st.sidebar.caption("Drop your CV here and let the pack work its magic. 🐶")

    uploaded = st.sidebar.file_uploader(
        "🐕 Fetch your CV (PDF, DOCX, or TXT)",
        type=["pdf", "docx", "txt", "md"],
    )

    if uploaded is not None:
        # Only re-extract if a new/different file was provided.
        if uploaded.name != st.session_state.cv_filename:
            try:
                text, fmt, raw = extract_text(uploaded)
            except ValueError as exc:
                st.sidebar.error(str(exc))
                return
            if not text:
                st.sidebar.error(
                    "🐱 Hmm, the kitties couldn't read any text in there. If it's "
                    "a scanned/image PDF, pop in a text-based version instead. 🐾"
                )
                return
            st.session_state.cv_text = text
            st.session_state.cv_bytes = raw
            st.session_state.cv_filename = uploaded.name
            st.session_state.cv_format = fmt
            # New CV invalidates any previous analysis/generation.
            st.session_state.advice = []
            st.session_state.analysis_done = False
            st.session_state.new_cv_text = ""
            st.session_state.new_cv_data = b""
            st.session_state.generation_done = False
            st.sidebar.success(
                f"🐶 Got it, {who()}! “{uploaded.name}” ({fmt.upper()}) is in good paws. 🐾"
            )

    if st.session_state.cv_filename:
        st.sidebar.divider()
        st.sidebar.markdown(f"**🐾 Your CV:** {st.session_state.cv_filename}")
        st.sidebar.markdown(
            f"**Format:** `{st.session_state.cv_format}`  ·  "
            f"**{len(st.session_state.cv_text):,}** characters sniffed out 🐕"
        )
        with st.sidebar.expander("👀 Peek at what we read"):
            st.text(st.session_state.cv_text[:5000])

    # --- Optional: who's using the app (personalises the greeting) --- #
    st.sidebar.divider()
    st.sidebar.text_input(
        "👋 What should we call you?",
        key="manual_name",
        placeholder="e.g. Dorka",
        help="Optional — personalises the app for you. If you're signed in to "
             "Streamlit, we may already know your name.",
    )


# --------------------------------------------------------------------------- #
# Main page
# --------------------------------------------------------------------------- #
def render_analysis() -> None:
    """Show the Recruiter & ATS findings: score, keywords, red flags."""
    score = st.session_state.match_score
    keywords = st.session_state.missing_keywords
    red_flags = st.session_state.red_flags

    st.subheader(f"🐶 The Recruiter Report, {who()} 🐱")
    st.caption(
        "Here's what a recruiter and the robot gatekeeper (ATS) see in your "
        "first 10 seconds. Every fix below is an easy win — you've got this! 💛"
    )

    # --- Match score --- #
    if score is not None:
        if score >= 75:
            mood = "Pawsitively impressive! The pack is wagging hard. 🐶🎉"
        elif score >= 50:
            mood = "A good sniff! A few tweaks and you'll be fetching interviews. 🐾"
        else:
            mood = "Rough start, but every great dog learns new tricks — let's climb! 🐕💪"
        c1, c2 = st.columns([0.32, 0.68])
        with c1:
            st.metric("🎯 Match Score", f"{score}/100")
        with c2:
            st.progress(score / 100)
            st.caption(mood)

    # --- Missing keywords --- #
    st.markdown("##### 🔑 Top Missing Keywords")
    if keywords:
        pills = " ".join(
            f"<span style='display:inline-block;background:#fde68a;color:#92400e;"
            f"padding:4px 12px;border-radius:999px;margin:3px;font-size:.9rem;"
            f"font-weight:600'>🦴 {k}</span>"
            for k in keywords
        )
        st.markdown(pills, unsafe_allow_html=True)
        st.caption("Sprinkle these into your CV (truthfully!) so the ATS sniffs you out. 🐾")
    else:
        st.success("🐶 Wow — no big keywords missing! Your CV already speaks their language.")

    # --- Red flags --- #
    st.markdown("##### 🚩 Quick-Spot Red Flags")
    st.caption("Things a busy hiring manager might trip over — all totally fixable below. 🐾")
    if red_flags:
        for flag in red_flags:
            st.warning(f"🐱 {flag}")
    else:
        st.success(f"🐾 No glaring red flags — nicely done, {who()}!")

    st.divider()


def render_advice() -> None:
    advice = st.session_state.advice
    if not advice:
        st.info("🐶 No specific rewrites this time — your bullets are already strong!")
        return

    accepted = sum(1 for a in advice if a["status"] == "accepted")
    rejected = sum(1 for a in advice if a["status"] == "rejected")
    pending = sum(1 for a in advice if a["status"] == "pending")

    st.subheader("✍️ Scroll-Stopping Rewrites From Your Career Pups 🐾")
    st.caption(
        "Each one uses the **Google XYZ formula** — *Accomplished X, as measured "
        f"by Y, by doing Z*. Keep the ones you love, {who()}! 🐶"
    )
    c1, c2, c3 = st.columns(3)
    c1.metric("🦴 Kept", accepted)
    c2.metric("🐈 Skipped", rejected)
    c3.metric("🐾 To review", pending)

    badge = {
        "pending": "🐾 Sniffing…",
        "accepted": "🦴 Kept!",
        "rejected": "🐈 Skipped",
    }
    kind_tag = {"xyz": "✍️ Bullet rewrite", "ats": "🤖 ATS / scroll-stopper"}

    for a in advice:
        with st.container(border=True):
            top, status_col = st.columns([0.75, 0.25])
            with top:
                tag = kind_tag.get(a.get("kind", "xyz"), "")
                st.markdown(f"**🐶 {a['title']}**  ·  *{tag}*")
                if a.get("original"):
                    st.markdown(
                        f"<span style='color:#9ca3af'>🚫 Before:</span> "
                        f"<span style='color:#6b7280'><s>{a['original']}</s></span>",
                        unsafe_allow_html=True,
                    )
                if a.get("suggestion"):
                    st.markdown(
                        f"<span style='color:#15803d;font-weight:600'>✨ After:</span> "
                        f"{a['suggestion']}",
                        unsafe_allow_html=True,
                    )
                if a.get("detail"):
                    st.caption(f"💡 {a['detail']}")
            with status_col:
                st.markdown(badge[a["status"]])

            b1, b2, _ = st.columns([0.25, 0.25, 0.5])
            with b1:
                if st.button(
                    "Keep it! 🦴",
                    key=f"accept_{a['id']}",
                    type="primary" if a["status"] != "accepted" else "secondary",
                    use_container_width=True,
                ):
                    a["status"] = "accepted"
                    st.rerun()
            with b2:
                if st.button(
                    "Skip 🐈",
                    key=f"reject_{a['id']}",
                    use_container_width=True,
                ):
                    a["status"] = "rejected"
                    st.rerun()


def render_interview_prep(safe_name: str) -> None:
    """Interview prep tab: generate tips on demand + one-click OneNote export."""
    st.markdown(
        "Walking into the interview? 🐾 Let the pack prep you with questions "
        "you'll likely get, your best talking points, and smart questions to ask."
    )

    if not st.session_state.interview_prep_done:
        if st.button(
            "🎤 Fetch my interview prep! 🐶",
            type="primary",
            use_container_width=True,
            key="prep_btn",
        ):
            client = get_client()
            if client is None:
                return
            with st.spinner(f"The corgis are coaching you for the big day, {who()}… 🐾🎤"):
                try:
                    st.session_state.interview_prep = generate_interview_prep(
                        client,
                        st.session_state.cv_text,
                        st.session_state.job_ad,
                        (st.session_state.company or "").strip(),
                    )
                    st.session_state.interview_prep_done = True
                    st.rerun()
                except Exception as exc:
                    show_api_error(exc, "preparing your interview tips")
        return

    prep = st.session_state.interview_prep
    sections = [
        ("likely_questions", "❓ Questions you'll likely be asked"),
        ("your_talking_points", "⭐ Your talking points"),
        ("questions_to_ask", "🙋 Smart questions to ask them"),
        ("quick_tips", "✅ Quick tips"),
    ]
    for key, label in sections:
        items = prep.get(key) or []
        if not items:
            continue
        st.markdown(f"**{label}**")
        for it in items:
            st.markdown(f"- {it}")

    st.divider()
    st.markdown("**📓 Add to OneNote** — one click, two easy ways:")
    title = (st.session_state.company or "").strip()
    title = f"Interview Prep — {title}" if title else "Interview Prep"
    prep_text = prep_to_text(prep, title)

    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "📓 Download for OneNote (.docx)",
            data=build_prep_docx(prep, title),
            file_name=f"{safe_name}_interview_prep.docx",
            mime=DOCX_MIME,
            type="primary",
            use_container_width=True,
        )
    with col2:
        st.link_button(
            "🔗 Open OneNote",
            "https://www.onenote.com/notebooks",
            use_container_width=True,
        )
    st.caption(
        "Easiest path: open OneNote → **Insert → File Printout** (or just drag "
        "the .docx onto a page). Prefer pasting? Copy the text below 👇"
    )
    st.code(prep_text, language="text")
    if st.button("🔄 Re-generate prep", key="prep_regen"):
        st.session_state.interview_prep_done = False
        st.session_state.interview_prep = {}
        st.rerun()


def render_generation() -> None:
    advice = st.session_state.advice
    accepted = [a for a in advice if a["status"] == "accepted"]
    reviewed = all(a["status"] != "pending" for a in advice)

    st.divider()
    st.subheader(f"🎁 Generate Tailored Documents, {who()}! 🐶🐱")

    fmt = st.session_state.cv_format
    style_preserved = fmt == "docx" and Document is not None

    company = (st.session_state.company or "").strip()
    style_note = (
        f"styled to match **{company}**'s vibe" if company
        else "in a clean, modern professional style"
    )
    st.markdown(
        "When you're happy with your choices, the pack will whip up **two** "
        "things just for this role: your **finalized résumé** *and* a matching "
        "**cover letter**. 🐾"
    )
    if style_preserved:
        st.caption(
            "🐶 Your **DOCX** résumé keeps its exact original styling — only the "
            "words are upgraded with your kept rewrites."
        )
    else:
        st.caption(
            f"🐾 Your résumé is **redesigned into a sleek, editable Word document "
            f"(.docx)** {style_note}, with your kept rewrites baked in — easy to "
            "tweak afterwards. 🐶"
        )

    if not reviewed:
        st.info(
            f"🐾 Almost there, {who()}! Give every rewrite a *Keep it!* or a *Skip* "
            "and the button below wakes up. 🐶"
        )

    if st.button(
        "🎁 Generate Tailored Documents 🐾",
        type="primary",
        disabled=not reviewed,
        use_container_width=True,
    ):
        client = get_client()
        if client is None:
            return
        with st.spinner(
            f"The corgis are writing your résumé AND cover letter, {who()}… 🐾🐶✍️"
        ):
            try:
                # --- 1. Finalized résumé --- #
                if style_preserved:
                    # DOCX upload → edit in place, keeping the user's styling.
                    data = regenerate_docx_in_place(
                        client,
                        st.session_state.cv_bytes,
                        st.session_state.job_ad,
                        accepted,
                        company,
                    )
                    preview = "\n".join(
                        p.text
                        for p in _iter_docx_paragraphs(Document(io.BytesIO(data)))
                    )
                else:
                    # PDF/TXT → redesign into a modern, editable, themed .docx.
                    try:
                        structured = regenerate_cv_structured(
                            client,
                            st.session_state.cv_text,
                            st.session_state.job_ad,
                            accepted,
                            company,
                        )
                        accent = _parse_accent(structured.get("theme"))
                        data = build_docx_styled(structured, accent)
                        preview = _structured_to_text(structured)
                    except Exception:
                        new_text = regenerate_cv(
                            client,
                            st.session_state.cv_text,
                            st.session_state.job_ad,
                            accepted,
                            company,
                        )
                        data = build_docx_plain(new_text)
                        preview = new_text
                mime = DOCX_MIME
                ext = "docx"

                # --- 2. Tailored cover letter (tone matches the résumé) --- #
                cover = generate_cover_letter(
                    client,
                    preview,
                    st.session_state.job_ad,
                    accepted,
                    company,
                )

                st.session_state.new_cv_data = data
                st.session_state.new_cv_mime = mime
                st.session_state.new_cv_ext = ext
                st.session_state.new_cv_text = preview
                st.session_state.cover_letter = cover
                st.session_state.generation_done = True
            except Exception as exc:
                show_api_error(exc, "writing your documents")
                return

    if st.session_state.generation_done:
        st.success(
            f"🐶 All done, {who()}! Your tailored **résumé** and **cover letter** "
            "are ready below. 🐾"
        )

        ext = st.session_state.new_cv_ext
        base = os.path.splitext(st.session_state.cv_filename or "cv")[0]
        safe = re.sub(r"[^A-Za-z0-9_-]+", "_", base).strip("_") or "cv"

        tab_resume, tab_cover, tab_prep = st.tabs(
            ["📄 Finalized Résumé", "✉️ Cover Letter", "🎤 Interview Prep"]
        )

        # --- Résumé tab --- #
        with tab_resume:
            st.markdown("**⬇️ Download your finalized résumé (Word, editable):**")
            st.download_button(
                f"🐶 Download Résumé (.{ext}) 🦴",
                data=st.session_state.new_cv_data,
                file_name=f"{safe}_tailored.{ext}",
                mime=st.session_state.new_cv_mime,
                type="primary",
                use_container_width=True,
            )
            with st.expander("👀 Peek at your résumé text"):
                st.text_area(
                    "Finalized résumé",
                    value=st.session_state.new_cv_text,
                    height=460,
                    label_visibility="collapsed",
                )
            if style_preserved:
                st.caption(
                    "Styling preserved from your original DOCX. (Formatting inside "
                    "a single paragraph may normalise to that paragraph's style.)"
                )
            else:
                st.caption(
                    "A clean, modern, fully editable Word document — open it and "
                    "fine-tune anything you like. ✨"
                )

        # --- Cover letter tab --- #
        with tab_cover:
            st.markdown(
                "**📋 Copy it** with the icon in the top-right of the box below, "
                "or **⬇️ download** it as a file. 🐾"
            )
            # st.code gives a one-click copy button in its corner.
            st.code(st.session_state.cover_letter, language="text")
            st.download_button(
                "✉️ Download Cover Letter (.txt) 🦴",
                data=st.session_state.cover_letter.encode("utf-8"),
                file_name=f"{safe}_cover_letter.txt",
                mime="text/plain",
                type="primary",
                use_container_width=True,
            )
            st.caption(
                "Tailored just for this role and tuned to match your résumé's "
                f"tone. Go get 'em, {who()}! 🐾🐱💛"
            )

        # --- Interview prep tab --- #
        with tab_prep:
            render_interview_prep(safe)


def render_main() -> None:
    st.title(f"{whos()} Career Magic 🐶🐱")
    st.markdown(
        f"#### Hey {who()}! 🐾 Job hunting is *exhausting* — but you've got this, "
        "and you've got a whole pack cheering you on. Let's make your CV "
        "irresistible together. 💛"
    )

    if not st.session_state.cv_filename:
        st.info(
            "🐶 **Ready when you are!** Pop your CV into the sidebar on the left "
            "and the corgis will get to work. 🐾"
        )
        with st.container():
            cols = st.columns([1, 1, 1])
            with cols[1]:
                show_corgi(width=260)
        return

    # --- Job advertisement input (separate text box) --- #
    st.subheader(f"🎯 The Job You Want, {who()}")
    st.caption("Paste the job ad below — the corgis will sniff out exactly what they're looking for. 🐕")
    st.session_state.job_ad = st.text_area(
        "Paste the job advertisement text",
        value=st.session_state.job_ad,
        height=240,
        placeholder=f"Paste the dream job ad here, {who()}… 🐾 We'll handle the rest!",
        label_visibility="collapsed",
    )

    # --- Optional company (drives the document's styling/colour) --- #
    st.session_state.company = st.text_input(
        "🏢 Company (optional)",
        value=st.session_state.company,
        placeholder="e.g. Spotify — leave empty for a general, neutral style",
        help="If given, the corgis colour your résumé to match the company's "
             "vibe. Empty = a clean, general professional look.",
    )

    if st.button(
        "🐶 Fetch My Tailoring Tips! 🐾", type="primary", use_container_width=True
    ):
        client = None
        if not st.session_state.job_ad.strip():
            st.warning(
                f"🐾 Oops — the corgis need a job ad to sniff first! Paste one above, {who()}. 🐶"
            )
        else:
            client = get_client()  # shows its own error if unavailable
        if client is not None:
            with st.spinner(
                f"The corgis are fetching the best keywords for you, {who()}… 🐾🐕"
            ):
                try:
                    result = analyze_cv(
                        client,
                        st.session_state.cv_text,
                        st.session_state.job_ad,
                    )
                    st.session_state.match_score = result["match_score"]
                    st.session_state.missing_keywords = result["missing_keywords"]
                    st.session_state.red_flags = result["red_flags"]
                    st.session_state.advice = result["recommendations"]
                    st.session_state.analysis_done = True
                    # A fresh analysis invalidates any previous documents.
                    st.session_state.generation_done = False
                    st.session_state.new_cv_text = ""
                    st.session_state.new_cv_data = b""
                    st.session_state.cover_letter = ""
                    st.session_state.interview_prep = {}
                    st.session_state.interview_prep_done = False
                except Exception as exc:
                    show_api_error(exc, "sniffing out your tailoring tips")

    if st.session_state.analysis_done:
        st.divider()
        render_analysis()
        render_advice()
        render_generation()


def main() -> None:
    # Resolve who's viewing once per run, before anything is rendered.
    st.session_state["_viewer_name"] = resolve_viewer_name()
    inject_theme()
    render_sidebar()
    render_main()


if __name__ == "__main__":
    main()
